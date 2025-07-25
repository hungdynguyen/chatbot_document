import os
import glob
import json
import pandas as pd
import docx
import fitz  # PyMuPDF
import pdfplumber
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from pathlib import Path
from typing import List, Optional
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_unstructured import UnstructuredLoader
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module='unstructured')

class DocumentParser:
    """
    A unified document parser for various file formats:
    - Excel (.xlsx, .xls)
    - PDF (.pdf)
    - Word (.docx, .doc)
    - Text (.txt)
    """
    
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 50, fallback_threshold: int = 5):
        """
        Initialize the document parser with configurable chunking parameters.
        
        Args:
            chunk_size: Maximum size of text chunks
            chunk_overlap: Overlap between consecutive chunks
        """
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        self.fallback_threshold = fallback_threshold  # Threshold for fallback logic
    

    def parse_file(self, file_path: str) -> List[Document]:
        """
        Parse file, ưu tiên parser chuyên biệt và sử dụng fallback nếu cần.
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        specialized_documents = []
        
        if file_extension in ['.xlsx', '.xls']:
            specialized_documents = self.parse_excel(file_path)
        elif file_extension == '.pdf':
            specialized_documents = self.parse_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            specialized_documents = self.parse_word(file_path)
        elif file_extension == '.txt':
            specialized_documents = self.parse_text(file_path)
        else:
            print(f"Unsupported file type: {file_extension}, sử dụng fallback.")
            return self._parse_with_unstructured_fallback(file_path)

        # --- LOGIC KIỂM TRA VÀ FALLBACK CỐT LÕI ---
        # Nếu parser chuyên biệt không trả về kết quả nào, hãy kích hoạt "lưới an toàn"
        if not specialized_documents:
            print(f"⚠️ Parser chuyên biệt cho file '{file_path.name}' không trích xuất được document nào.")
            return self._parse_with_unstructured_fallback(file_path)
        if 0 < len(specialized_documents) < self.fallback_threshold:
            print(f"⚠️ Số lượng document trích xuất được từ file '{file_path.name}' dưới ngưỡng tối thiểu.")
            fallback_docs = self._parse_with_unstructured_fallback(file_path)
            return specialized_documents + fallback_docs
        print(f"✅ Parser chuyên biệt cho file '{file_path.name}' đã hoạt động thành công.")
        return specialized_documents
    
    def parse_directory(self, directory_path: str, output_dir: Optional[str] = None) -> dict:
        """
        Parse all supported files in a directory and save results to JSON files.
        
        Args:
            directory_path: Path to the directory containing files to parse
            output_dir: Directory to save the output JSON files (optional)
            
        Returns:
            Dictionary with counts of parsed documents by file type
        """
        directory_path = Path(directory_path)
        
        # Define output directory
        if output_dir:
            output_dir = Path(output_dir)
        else:
            output_dir = directory_path / "output"
        
        # Create output directory if it doesn't exist
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Find files by extension
        excel_files = list(directory_path.glob("*.xlsx")) + list(directory_path.glob("*.xls"))
        pdf_files = list(directory_path.glob("*.pdf"))
        word_files = list(directory_path.glob("*.docx")) + list(directory_path.glob("*.doc"))
        text_files = list(directory_path.glob("*.txt"))
        
        # Print summary of files found
        print(f"Found {len(excel_files)} Excel files, {len(pdf_files)} PDF files, "
              f"{len(word_files)} Word files, and {len(text_files)} text files.")
        
        # Dictionary to store counts
        results = {
            "excel": 0,
            "pdf": 0,
            "word": 0,
            "text": 0
        }
        
        # Process Excel files
        if excel_files:
            all_excel_documents = []
            for file_path in excel_files:
                print(f"Parsing Excel file: {file_path}")
                documents = self.parse_excel(file_path)
                all_excel_documents.extend(documents)
                print(f"Parsed {len(documents)} documents from {file_path}")
            
            # Save Excel results
            if all_excel_documents:
                excel_output = output_dir / "excel_documents_parsed.json"
                with excel_output.open("w", encoding="utf-8") as f:
                    json.dump([{"page_content": doc.page_content, "metadata": doc.metadata} 
                              for doc in all_excel_documents], f, ensure_ascii=False, indent=4)
                results["excel"] = len(all_excel_documents)
        
        # Process PDF files
        if pdf_files:
            all_pdf_documents = []
            for file_path in pdf_files:
                print(f"Parsing PDF file: {file_path}")
                documents = self.parse_pdf(file_path)
                all_pdf_documents.extend(documents)
                print(f"Parsed {len(documents)} documents from {file_path}")
            
            # Save PDF results
            if all_pdf_documents:
                pdf_output = output_dir / "pdf_documents_parsed.json"
                with pdf_output.open("w", encoding="utf-8") as f:
                    json.dump([{"page_content": doc.page_content, "metadata": doc.metadata} 
                              for doc in all_pdf_documents], f, ensure_ascii=False, indent=4)
                results["pdf"] = len(all_pdf_documents)
        
        # Process Word files
        if word_files:
            all_word_documents = []
            for file_path in word_files:
                print(f"Parsing Word file: {file_path}")
                documents = self.parse_word(file_path)
                all_word_documents.extend(documents)
                print(f"Parsed {len(documents)} documents from {file_path}")
            
            # Save Word results
            if all_word_documents:
                word_output = output_dir / "word_documents_parsed.json"
                with word_output.open("w", encoding="utf-8") as f:
                    json.dump([{"page_content": doc.page_content, "metadata": doc.metadata} 
                              for doc in all_word_documents], f, ensure_ascii=False, indent=4)
                results["word"] = len(all_word_documents)
        
        # Process Text files
        if text_files:
            all_text_documents = []
            for file_path in text_files:
                print(f"Parsing text file: {file_path}")
                documents = self.parse_text(file_path)
                all_text_documents.extend(documents)
                print(f"Parsed {len(documents)} documents from {file_path}")
            
            # Save Text results
            if all_text_documents:
                text_output = output_dir / "text_documents_parsed.json"
                with text_output.open("w", encoding="utf-8") as f:
                    json.dump([{"page_content": doc.page_content, "metadata": doc.metadata} 
                              for doc in all_text_documents], f, ensure_ascii=False, indent=4)
                results["text"] = len(all_text_documents)
        
        # Save combined results
        all_documents = []
        for file_type in ["excel", "pdf", "word", "text"]:
            input_file = output_dir / f"{file_type}_documents_parsed.json"
            if input_file.exists():
                with input_file.open("r", encoding="utf-8") as f:
                    all_documents.extend(json.load(f))
        
        if all_documents:
            combined_output = output_dir / "all_documents_parsed.json"
            with combined_output.open("w", encoding="utf-8") as f:
                json.dump(all_documents, f, ensure_ascii=False, indent=4)
        
        print(f"Total documents created: {sum(results.values())}")
        return results
    
    
    # Thêm phương thức này vào bên trong class DocumentParser

    def _parse_with_unstructured_fallback(self, file_path: Path) -> List[Document]:
        """
        Phương thức dự phòng, sử dụng UnstructuredLoader để đảm bảo không bỏ sót dữ liệu.
        """
        print(f"    -> [Fallback] Sử dụng UnstructuredLoader cho file {file_path.name}...")
        try:
            # Dùng UnstructuredLoader để hút toàn bộ text thô
            loader = UnstructuredLoader(str(file_path))
            raw_documents = loader.load()
            
            # Gộp tất cả thành một khối và chia lại bằng text_splitter của class
            full_text = "\n\n".join([doc.page_content for doc in raw_documents])
            if not full_text.strip():
                return []
            
            chunks = self.text_splitter.create_documents([full_text])
            
            # Quan trọng: Thêm metadata chuẩn hóa cho từng chunk
            for chunk in chunks:
                chunk.metadata = {
                    "source": str(file_path),
                    "basename": file_path.name,
                    "file_type": file_path.suffix.lower(),
                    "content_type": "unstructured_chunk",
                    "parser_method": "fallback"
                }
            print(f"    -> [Fallback] Tạo thành công {len(chunks)} chunks.")
            return chunks
        except Exception as e:
            print(f"    -> [Fallback] Lỗi khi chạy UnstructuredLoader: {e}")
            return []
    
    def parse_excel(self, file_path: Path) -> List[Document]:
        documents = []
        try:
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                if df.empty:
                    continue
                
                # --- Logic tìm bảng của bạn vẫn giữ nguyên ---
                df['is_empty'] = df.isnull().all(axis=1)
                df['table_id'] = df['is_empty'].cumsum()
                tables = df.groupby('table_id')
                
                for table_id, table_df in tables:
                    table_df = table_df.dropna(how='all').reset_index(drop=True)
                    table_df = table_df.drop(columns=['is_empty', 'table_id'], errors='ignore')
                    if table_df.empty:
                        continue
                    
                    # --- BẮT ĐẦU CẢI TIẾN ---

                    # Cải tiến 1: Tạo một Document cho TOÀN BỘ bảng dưới dạng Markdown
                    # Điều này giữ lại ngữ cảnh so sánh giữa các hàng.
                    markdown_text = table_df.to_markdown(index=False)
                    table_content = f"Bảng dữ liệu từ sheet '{sheet_name}':\n\n{markdown_text}"
                    table_metadata = {
                        "source": str(file_path), "file_type": "excel", "sheet_name": sheet_name,
                        "table_id": f"table_{table_id}", "content_type": "full_table_markdown"
                    }
                    documents.append(Document(page_content=table_content, metadata=table_metadata))

                    # Cải tiến 2: Vẫn giữ lại việc parse từng hàng để truy xuất chi tiết
                    headers = table_df.columns.tolist()
                    for index, row in table_df.iterrows():
                        row_texts = [f"{str(h).strip()}: {str(row[h]).strip()}" for h in headers if pd.notna(row[h])]
                        if not row_texts: continue
                        
                        row_content = " | ".join(row_texts)
                        row_metadata = {
                            "source": str(file_path), "file_type": "excel", "sheet_name": sheet_name,
                            "table_id": f"table_{table_id}", "row_index_in_table": index,
                            "content_type": "table_row"
                        }
                        documents.append(Document(page_content=row_content, metadata=row_metadata))

                # Add UnstructuredLoader
                
            loader = UnstructuredLoader(str(file_path))
            unstructured_docs = loader.load()
            # Thêm metadata để phân biệt
            for doc in unstructured_docs:
                doc.metadata.update({
                    "source": str(file_path),
                    "file_type": "excel",
                    "content_type": "unstructured_chunk",
                    "parser_method": "unstructured"
                })
            documents.extend(unstructured_docs)

        
        except Exception as e:
            print(f"Error parsing Excel file with UnstructuredLoader {file_path}: {e}")
        return documents
    
    
    # def parse_excel(self, file_path: Path) -> List[Document]:
    #     documents = []
    #     try:
    #         excel_file = pd.ExcelFile(file_path)
            
    #         for sheet_name in excel_file.sheet_names:
    #             df = pd.read_excel(excel_file, sheet_name=sheet_name)
    #             if df.empty:
    #                 continue
                    
    #             df['is_empty'] = df.isnull().all(axis=1)
    #             df['table_id'] = (~df['is_empty']).cumsum().where(df['is_empty']).ffill().fillna(0)
                
    #             tables = df[~df['is_empty']].groupby('table_id')
                
    #             for table_id, table_df in tables:
    #                 table_df = table_df.drop(columns=['is_empty', 'table_id'], errors='ignore').dropna(how='all')
    #                 if table_df.empty or len(table_df) < 2: # Bỏ qua các bảng quá nhỏ
    #                     continue
                    
    #                 # --- THAY ĐỔI CỐT LÕI NẰM Ở ĐÂY ---
    #                 # Chuyển đổi toàn bộ bảng thành văn bản Markdown
    #                 markdown_text = table_df.to_markdown(index=False)
                    
    #                 # Tạo một Document duy nhất cho mỗi bảng nhỏ thay vì mỗi hàng
    #                 page_content = f"Bảng dữ liệu từ sheet '{sheet_name}':\n\n{markdown_text}"
                    
    #                 metadata = {
    #                     "source": str(file_path),
    #                     "file_type": "excel",
    #                     "sheet_name": sheet_name,
    #                     "table_id": f"table_{int(table_id)}",
    #                     "content_type": "markdown_table" # Đánh dấu đây là một bảng
    #                 }
                    
    #                 # Thay vì chia nhỏ bảng, ta coi cả bảng là một chunk duy nhất
    #                 # Nếu bảng quá lớn, có thể chia nó thành các chunk nhỏ hơn sau này
    #                 documents.append(Document(
    #                     page_content=page_content,
    #                     metadata=metadata
    #                 ))

    #     except Exception as e:
    #         print(f"Error parsing Excel file (optimized with Markdown) {file_path}: {e}")
                
    #     return documents
    
    def parse_pdf(self, file_path: Path) -> List[Document]:
        """
        Optimized PDF parser: extracts tables row-by-row and text paragraph-by-paragraph.
        Includes OCR fallback.
        """
        documents = []
        
        # --- Method 1: Direct text and table extraction ---
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    # 1. Extract tables first
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        # Giả định hàng đầu tiên là header
                        headers = [h.strip() if h else f"col_{i}" for i, h in enumerate(table[0])]
                        for row_idx, row in enumerate(table[1:], start=1):
                            row_texts = [
                                f"{headers[i]}: {str(cell).strip()}"
                                for i, cell in enumerate(row) if cell and str(cell).strip()
                            ]
                            if not row_texts: continue
                            
                            page_content = " | ".join(row_texts)
                            metadata = {
                                "source": str(file_path), "file_type": "pdf", "page_num": page_num,
                                "content_type": "table_row", "table_id": table_idx, 
                                "row_index_in_table": row_idx, "extraction_method": "text"
                            }
                            documents.append(Document(page_content=page_content, metadata=metadata))

                    # 2. Extract text and chunk it
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_chunks = self.text_splitter.split_text(page_text)
                        for chunk in text_chunks:
                            metadata = {
                                "source": str(file_path), "file_type": "pdf", "page_num": page_num,
                                "content_type": "paragraph", "extraction_method": "text"
                            }
                            documents.append(Document(page_content=chunk, metadata=metadata))
                            
        except Exception as e:
            print(f"Error during direct PDF parsing for {file_path}: {e}")

        # --- Method 2: OCR Fallback if direct extraction yields little content ---
        if not documents or len("".join(d.page_content for d in documents)) < 100:
            print(f"PDF {file_path} has little text, trying OCR...")
            try:
                images = convert_from_path(file_path, dpi=300)
                for page_num, image in enumerate(images, start=1):
                    ocr_text = pytesseract.image_to_string(image, lang='vie+eng')
                    if ocr_text and ocr_text.strip():
                        text_chunks = self.text_splitter.split_text(ocr_text)
                        for chunk in text_chunks:
                            metadata = {
                                "source": str(file_path), "file_type": "pdf", "page_num": page_num,
                                "content_type": "paragraph", "extraction_method": "ocr"
                            }
                            documents.append(Document(page_content=chunk, metadata=metadata))
            except Exception as e:
                print(f"Error during OCR PDF parsing for {file_path}: {e}")

        return documents
    
    
    
    def parse_word(self, file_path: Path) -> List[Document]:
        """
        Hybrid Word parser: extracts paragraphs, full tables as Markdown, and individual table rows.
        This approach provides both contextual (full table) and specific (single row) data for the AI.
        """
        documents = []
        try:
            doc = docx.Document(file_path)

            # --- 1. Process tables using the hybrid strategy ---
            for table_idx, table in enumerate(doc.tables):
                # First, read all table data into a list of lists to be reused.
                # This is more efficient than iterating through the table multiple times.
                table_data = []
                for row in table.rows:
                    # Skip fully empty rows to avoid noise
                    if not any(cell.text.strip() for cell in row.cells):
                        continue
                    table_data.append([cell.text.strip() for cell in row.cells])
                
                # If the table has no meaningful data, skip it
                if not table_data:
                    continue

                # --- Strategy A: Create one Document for the ENTIRE table as Markdown ---
                # This preserves the overall context of the table for comparison/summary tasks.
                try:
                    headers = table_data[0]
                    markdown_table = "| " + " | ".join(str(h) for h in headers) + " |\n"
                    markdown_table += "| " + " | ".join(['---'] * len(headers)) + " |\n"
                    for row_cells in table_data[1:]:
                        # Pad the row to match header length, preventing errors
                        padded_row = row_cells + [''] * (len(headers) - len(row_cells))
                        markdown_table += "| " + " | ".join(str(cell) for cell in padded_row) + " |\n"

                    markdown_metadata = {
                        "source": str(file_path),
                        "file_type": "word",
                        "content_type": "full_table_markdown", # Descriptive type
                        "table_id": table_idx
                    }
                    documents.append(Document(page_content=markdown_table, metadata=markdown_metadata))
                except IndexError as e:
                    print(f"    - Warning: Could not process table {table_idx} as Markdown in {file_path.name}: {e}")


                # --- Strategy B: Create a Document for EACH ROW in the table ---
                # This is excellent for precise, specific data retrieval.
                headers = table_data[0]
                for row_idx, row_cells in enumerate(table_data[1:], start=1):
                    row_texts = []
                    for i, cell_text in enumerate(row_cells):
                        if cell_text:
                            # Use header if available, otherwise use a generic column name
                            header = headers[i] if i < len(headers) else f"col_{i}"
                            row_texts.append(f"{header}: {cell_text}")
                    
                    if not row_texts: continue

                    row_metadata = {
                        "source": str(file_path),
                        "file_type": "word",
                        "content_type": "table_row", # Descriptive type
                        "table_id": table_idx,
                        "row_index_in_table": row_idx
                    }
                    documents.append(Document(page_content=" | ".join(row_texts), metadata=row_metadata))

            # --- 2. Extract and chunk paragraph text (as before) ---
            # The doc.paragraphs object intelligently excludes text that is inside tables.
            full_text = "\n\n".join(
                para.text.strip() for para in doc.paragraphs if para.text.strip()
            )
            
            if full_text:
                text_chunks = self.text_splitter.split_text(full_text)
                for chunk in text_chunks:
                    metadata = {
                        "source": str(file_path),
                        "file_type": "word",
                        "content_type": "paragraph"
                    }
                    documents.append(Document(page_content=chunk, metadata=metadata))

        except Exception as e:
            print(f"Error parsing Word file (hybrid) {file_path}: {e}")

        return documents
    
    # def parse_word(self, file_path: Path) -> List[Document]:
    #     """
    #     Optimized Word parser: extracts tables row-by-row and text paragraph-by-paragraph.
    #     """
    #     documents = []
    #     try:
    #         doc = docx.Document(file_path)

    #         # 1. Extract tables first
    #         for table_idx, table in enumerate(doc.tables):
    #             if not table.rows:
    #                 continue
                
    #             # Assume the first row is the header
    #             headers = [cell.text.strip() for cell in table.rows[0].cells]
                
    #             # Iterate over data rows
    #             for row_idx, row in enumerate(table.rows[1:], start=1):
    #                 row_texts = []
    #                 for i, cell in enumerate(row.cells):
    #                     cell_text = cell.text.strip()
    #                     if cell_text:
    #                         # Use header if available, otherwise use column index
    #                         header = headers[i] if i < len(headers) else f"col_{i}"
    #                         row_texts.append(f"{header}: {cell_text}")
                    
    #                 if not row_texts: continue

    #                 page_content = " | ".join(row_texts)
    #                 metadata = {
    #                     "source": str(file_path),
    #                     "file_type": "word",
    #                     "content_type": "table_row",
    #                     "table_id": table_idx,
    #                     "row_index_in_table": row_idx
    #                 }
    #                 documents.append(Document(page_content=page_content, metadata=metadata))
            
    #         # 2. Extract and chunk paragraph text
    #         # The doc.paragraphs object intelligently excludes text within tables.
    #         full_text = "\n\n".join(
    #             para.text.strip() for para in doc.paragraphs if para.text.strip()
    #         )
            
    #         if full_text:
    #             text_chunks = self.text_splitter.split_text(full_text)
    #             for chunk in text_chunks:
    #                 metadata = {
    #                     "source": str(file_path),
    #                     "file_type": "word",
    #                     "content_type": "paragraph"
    #                 }
    #                 documents.append(Document(page_content=chunk, metadata=metadata))

    #     except Exception as e:
    #         print(f"Error parsing Word file (optimized) {file_path}: {e}")

    #     return documents
    
    
    
    # def parse_word(self, file_path: Path) -> List[Document]:
    #     """
    #     Optimized Word parser: converts tables to Markdown and extracts paragraphs.
    #     """
    #     documents = []
    #     try:
    #         doc = docx.Document(file_path)

    #         # 1. Extract tables as single Markdown chunks
    #         for table_idx, table in enumerate(doc.tables):
    #             # Read table data into a list of lists
    #             table_data = []
    #             for row in table.rows:
    #                 # Bỏ qua các hàng trống hoàn toàn
    #                 if not any(cell.text.strip() for cell in row.cells):
    #                     continue
    #                 table_data.append([cell.text.strip() for cell in row.cells])
                
    #             if not table_data:
    #                 continue
                
    #             # Convert the list of lists to a Markdown table string
    #             # Header
    #             markdown_table = "| " + " | ".join(str(header) for header in table_data[0]) + " |\n"
    #             # Separator
    #             markdown_table += "| " + " | ".join(['---'] * len(table_data[0])) + " |\n"
    #             # Body
    #             for row in table_data[1:]:
    #                 markdown_table += "| " + " | ".join(str(cell) for cell in row) + " |\n"

    #             page_content = markdown_table
    #             metadata = {
    #                 "source": str(file_path),
    #                 "basename": file_path.name,
    #                 "file_type": "word",
    #                 "content_type": "markdown_table",
    #                 "table_id": table_idx,
    #             }
    #             documents.append(Document(page_content=page_content, metadata=metadata))
            
    #         # 2. Extract and chunk paragraph text
    #         # The doc.paragraphs object intelligently excludes text within tables.
    #         full_text = "\n\n".join(
    #             para.text.strip() for para in doc.paragraphs if para.text.strip()
    #         )
            
    #         if full_text:
    #             text_chunks = self.text_splitter.split_text(full_text)
    #             for chunk in text_chunks:
    #                 metadata = {
    #                     "source": str(file_path),
    #                     "basename": file_path.name,
    #                     "file_type": "word",
    #                     "content_type": "paragraph"
    #                 }
    #                 documents.append(Document(page_content=chunk, metadata=metadata))

    #     except Exception as e:
    #         print(f"Error parsing Word file (upgraded markdown version) {file_path}: {e}")

    #     return documents
    
    def parse_text(self, file_path: Path) -> List[Document]:
        """
        Optimized Text parser: reads the entire file and splits it into manageable
        text chunks using the class's text_splitter.
        """
        documents = []
        content = ""
        encoding_used = "utf-8"
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            print(f"UTF-8 decoding failed for {file_path}. Trying latin-1.")
            encoding_used = "latin-1"
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
            except Exception as e:
                print(f"Error reading text file {file_path} with latin-1: {e}")
                return []
        except Exception as e:
            print(f"Error reading text file {file_path}: {e}")
            return []

        if content.strip():
            text_chunks = self.text_splitter.split_text(content)
            for i, chunk in enumerate(text_chunks):
                metadata = {
                    "source": str(file_path),
                    "file_type": "text",
                    "content_type": "text_chunk",
                    "chunk_id": i,
                    "encoding": encoding_used
                }
                documents.append(Document(page_content=chunk, metadata=metadata))
        
        return documents


# Example usage
if __name__ == "__main__":
    # Initialize parser
    parser = DocumentParser()
    
    # Example 1: Parse a single file
    # file_path = "/path/to/document.pdf"
    # documents = parser.parse_file(file_path)
    # print(f"Parsed {len(documents)} documents from {file_path}")
    
    # Example 2: Parse all supported files in a directory
    data_dir = "/mnt/d/Techcombank_/chatbot_document/data/data_real"
    output_dir = "/mnt/d/Techcombank_/chatbot_document/data/output"
    results = parser.parse_directory(data_dir, output_dir)
    
    print("\nParsing Summary:")
    print(f"Excel documents: {results['excel']}")
    print(f"PDF documents: {results['pdf']}")
    print(f"Word documents: {results['word']}")
    print(f"Text documents: {results['text']}")
    print(f"Total documents: {sum(results.values())}")
