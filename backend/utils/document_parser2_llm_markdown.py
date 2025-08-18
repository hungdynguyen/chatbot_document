import os
import json
import time
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional
import google.generativeai as genai
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import warnings
from config import GOOGLE_API_KEY, TOGETHER_API_KEY
from langchain_together import ChatTogether
from langchain_core.messages import HumanMessage
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
import base64
import textwrap
warnings.filterwarnings("ignore", category=UserWarning, module='unstructured')

class DocumentParser:
    """
    Document Parser với conversion support nâng cao cho Gemini.
    Tích hợp file_converter để hỗ trợ nhiều format conversion.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 150):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        
        # Rate limiting settings
        self.last_api_call = 0
        self.min_delay_between_calls = 10.0  # Tăng từ 5 lên 10 giây
        self.max_retries = 5  # Tăng số lần thử lại
        self.base_retry_delay = 15.0  # Tăng thời gian chờ cơ bản
        self.exponential_backoff = True  # Thêm backoff theo cấp số nhân
        
        # Thêm giới hạn số lượng request mỗi phút
        self.max_requests_per_minute = 30
        self.request_count = 0
        self.minute_window_start = time.time()
        
        # API Setup
        if not GOOGLE_API_KEY:
            raise ValueError("GOOGLE_API_KEY không được tìm thấy trong config.py")
        
        genai.configure(api_key=GOOGLE_API_KEY)
        
        try:
            self.llm_client = genai.GenerativeModel('gemini-2.0-flash')
            print("✅ Gemini client (gemini-2.0-flash) for parsing đã sẵn sàng.")
        except Exception as e:
            print(f"⚠️ Fallback to gemini-1.5-flash: {e}")
            self.llm_client = genai.GenerativeModel('gemini-1.5-flash')
        
        
        # if not TOGETHER_API_KEY:
        #     raise ValueError("TOGETHER_API_KEY không được tìm thấy trong config.py")

        # try:
        #     self.llm_client = ChatTogether(
        #         model="meta-llama/Llama-3.2-90B-Vision-Instruct-Turbo",
        #         temperature=0,
        #         together_api_key=TOGETHER_API_KEY
        #     )
        #     print(f"✅ Together AI client Llama3.2-90B for parsing đã sẵn sàng.")
        # except Exception as e:
        #     print(f"❌ Lỗi khi khởi tạo Together AI client: {e}")
        #     raise

        # Supported file types cho direct processing
        self.supported_direct_types = {".pdf", ".txt", ".png", ".jpg", ".jpeg", ".gif", ".webp"}
        self.convertible_types = {".docx", ".doc", ".xlsx", ".xls"}
        
        # Check system capabilities
        self.libreoffice_available = self._check_libreoffice()
        self.excel_pdf_available = self._check_excel_pdf_libs()
        
        print(f"🔧 System capabilities:")
        print(f"   - LibreOffice: {'✅' if self.libreoffice_available else '❌'}")
        print(f"   - Excel-to-PDF: {'✅' if self.excel_pdf_available else '❌'}")

    def _check_libreoffice(self) -> bool:
        """Check if LibreOffice is available for document conversion."""
        try:
            result = subprocess.run(['libreoffice', '--version'], 
                                  capture_output=True, text=True, timeout=5)
            return result.returncode == 0
        except (subprocess.TimeoutExpired, FileNotFoundError):
            return False

    def _check_excel_pdf_libs(self) -> bool:
        """Check if required libraries for Excel-to-PDF conversion are available."""
        try:
            import pandas as pd
            from reportlab.lib.pagesizes import A4
            return True
        except ImportError:
            return False

    def _convert_to_supported_format(self, file_path: Path) -> Path:
        """
        Convert unsupported files to supported format using enhanced converters.
        """
        file_extension = file_path.suffix.lower()
        
        if file_extension in [".docx", ".doc"]:
            return self._convert_word_to_pdf_enhanced(file_path)
        elif file_extension in [".xlsx", ".xls"]:
            return self._convert_excel_to_pdf_enhanced(file_path)
        else:
            raise ValueError(f"Conversion not supported for {file_extension}")

    def _convert_word_to_pdf_enhanced(self, word_path: Path) -> Path:
        """
        Enhanced Word to PDF conversion with LibreOffice support.
        """
        print(f"🔄 Converting Word to PDF: {word_path.name}")
        
        # Create temporary PDF file
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_file:
            pdf_path = Path(tmp_file.name)
        
        # Method 1: Try LibreOffice (preferred for Linux)
        if self.libreoffice_available:
            try:
                cmd = [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(pdf_path.parent),
                    str(word_path)
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    # LibreOffice creates PDF with same name as input file
                    generated_pdf = pdf_path.parent / f"{word_path.stem}.pdf"
                    if generated_pdf.exists():
                        if generated_pdf != pdf_path:
                            generated_pdf.rename(pdf_path)
                        print(f"✅ LibreOffice conversion successful: {pdf_path}")
                        return pdf_path
                    
            except subprocess.TimeoutExpired:
                print("⚠️ LibreOffice conversion timed out")
            except Exception as e:
                print(f"⚠️ LibreOffice conversion failed: {e}")
        
        # Method 2: Fallback to text extraction
        print("🔄 Falling back to text extraction...")
        return self._extract_docx_text_to_temp_file(word_path)

    def _convert_excel_to_pdf_enhanced(self, excel_path: Path) -> Path:
        """
        Enhanced Excel to PDF conversion maintaining table structure.
        """
        print(f"🔄 Converting Excel to PDF: {excel_path.name}")
        
        if self.excel_pdf_available:
            try:
                # Create temporary PDF file
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_file:
                    pdf_path = Path(tmp_file.name)
                
                # Read Excel file
                df_dict = pd.read_excel(excel_path, sheet_name=None)  # Read all sheets
                
                # Create PDF document
                doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
                elements = []
                styles = getSampleStyleSheet()
                
                for sheet_name, sheet_df in df_dict.items():
                    # Add sheet title
                    title = Paragraph(f"<b>{sheet_name}</b>", styles['Heading1'])
                    elements.append(title)
                    elements.append(Spacer(1, 12))
                    
                    # Convert dataframe to table data
                    if not sheet_df.empty:
                        table_data = [list(sheet_df.columns)]  # Headers
                        for row in sheet_df.values:
                            table_data.append([str(cell) if pd.notna(cell) else '' for cell in row])
                        
                        # Create table with styling
                        table = Table(table_data)
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 8),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        
                        elements.append(table)
                        elements.append(Spacer(1, 20))
                
                # Build PDF
                doc.build(elements)
                print(f"✅ Excel-to-PDF conversion successful: {pdf_path}")
                return pdf_path
                
            except Exception as e:
                print(f"⚠️ Excel-to-PDF conversion failed: {e}")
        
        # Fallback to text extraction
        print("🔄 Falling back to text extraction...")
        return self._extract_xlsx_text_to_temp_file(excel_path)

    def _extract_docx_text_to_temp_file(self, docx_path: Path) -> Path:
        """
        Extract text from DOCX and save to temporary text file.
        """
        print(f"🔄 Extracting text from DOCX: {docx_path.name}")
        
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(docx_path)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    full_text.extend(table_text)
                    full_text.append("")  # Empty line after table
            
            # Save to temporary text file
            with tempfile.NamedTemporaryFile(mode='w', suffix=".txt", delete=False, encoding='utf-8') as tmp_file:
                tmp_file.write('\n'.join(full_text))
                temp_path = Path(tmp_file.name)
            
            print(f"✅ Extracted to text file: {temp_path}")
            return temp_path
            
        except Exception as e:
            print(f"❌ Text extraction failed: {e}")
            raise

    def _extract_xlsx_text_to_temp_file(self, xlsx_path: Path) -> Path:
        """
        Extract data from XLSX and save to temporary text file.
        """
        print(f"🔄 Extracting data from XLSX: {xlsx_path.name}")
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(xlsx_path)
            full_text = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                if not df.empty:
                    full_text.append(f"=== SHEET: {sheet_name} ===")
                    
                    # Convert to markdown-style table
                    headers = " | ".join(str(col) for col in df.columns)
                    separator = " | ".join("---" for _ in df.columns)
                    full_text.append(headers)
                    full_text.append(separator)
                    
                    for _, row in df.iterrows():
                        row_text = " | ".join(str(cell) if pd.notna(cell) else "" for cell in row)
                        full_text.append(row_text)
                    
                    full_text.append("")  # Empty line between sheets
            
            # Save to temporary text file
            with tempfile.NamedTemporaryFile(mode='w', suffix=".txt", delete=False, encoding='utf-8') as tmp_file:
                tmp_file.write('\n'.join(full_text))
                temp_path = Path(tmp_file.name)
            
            print(f"✅ Extracted to text file: {temp_path}")
            return temp_path
            
        except Exception as e:
            print(f"❌ XLSX extraction failed: {e}")
            raise
    
    def _apply_rate_limit(self):
        """Enhanced rate limiting with multiple strategies."""
        current_time = time.time()
        
        # 1. Basic delay between calls
        time_since_last_call = current_time - self.last_api_call
        if time_since_last_call < self.min_delay_between_calls:
            delay = self.min_delay_between_calls - time_since_last_call
            print(f"⏳ Basic rate limiting: Waiting {delay:.1f} seconds...")
            time.sleep(delay)
        
        # 2. Requests per minute limit
        if current_time - self.minute_window_start > 60:
            # Reset counter if we're in a new minute window
            self.request_count = 0
            self.minute_window_start = current_time
        else:
            self.request_count += 1
        
        if self.request_count >= self.max_requests_per_minute:
            time_remaining = 60 - (current_time - self.minute_window_start)
            print(f"⏳ Minute limit reached. Waiting {time_remaining:.1f} seconds...")
            time.sleep(time_remaining)
            self.request_count = 0
            self.minute_window_start = time.time()
        
        self.last_api_call = time.time()

    def _parse_with_llm_with_retry(self, file_path: Path) -> str:
        """
        Parse file với enhanced conversion support và retry logic.
        """
        original_path = file_path
        temp_file_created = False
        
        # Check if conversion is needed
        if file_path.suffix.lower() not in self.supported_direct_types:
            if file_path.suffix.lower() in self.convertible_types:
                print(f"📋 File type {file_path.suffix} không được Gemini hỗ trợ trực tiếp. Đang convert...")
                try:
                    file_path = self._convert_to_supported_format(file_path)
                    temp_file_created = True
                except Exception as conv_error:
                    print(f"❌ Conversion failed: {conv_error}")
                    return ""
            else:
                print(f"❌ File type {file_path.suffix} không được hỗ trợ.")
                return ""

        print(f"🧠 Bắt đầu parsing bằng LLM cho file: {original_path.name}...")
        
        uploaded_file = None
        
        try:
            for attempt in range(self.max_retries):
                try:
                    # Rate limiting
                    self._apply_rate_limit()
                    
                    # Upload file
                    if not uploaded_file:
                        uploaded_file = genai.upload_file(path=str(file_path), display_name=original_path.name)
                    
                    # Wait for processing
                    max_wait_time = 60
                    wait_time = 0
                    while uploaded_file.state.name == "PROCESSING" and wait_time < max_wait_time:
                        time.sleep(2)
                        wait_time += 2
                        uploaded_file = genai.get_file(uploaded_file.name)
                    
                    if uploaded_file.state.name == "FAILED":
                        print(f"❌ File upload failed: {uploaded_file.state}")
                        return ""
                    
                    # Generate content
                    self._apply_rate_limit()
                    
                    prompt = """
                    <TASK_DEFINITION>
                    You are an automated data processing engine. Your sole task is to analyze the provided file and convert its entire content into a single, clean, and well-structured Markdown string. The output must be a perfect representation of the original data, suitable for machine parsing later.

                    Follow these critical instructions precisely:

                    1.  **Analyze Layout:** First, analyze the visual layout of the document. Identify key-value pairs (e.g., a label in one cell and its value in another, potentially non-adjacent cell) and structured tables.
                    2.  **Convert Key-Value Pairs:** Represent all identified key-value pairs clearly.
                    3.  **Convert Tables:** Convert all structured tables into standard Markdown table format.
                    4.  **Preserve Content:** All text and numerical data must be preserved exactly as it appears in the original file.
                    5.  **No Extra Content:** Do not add any summaries, explanations, comments, or any text that is not present in the original document.
                    6.  **Strict Output Format:** Your entire output must be ONLY the Markdown content. Do not wrap it in ```markdown ... ``` or any other formatting.
                    </TASK_DEFINITION>

                    <OUTPUT_EXAMPLE>
                    ### A. THÔNG TIN CHUNG
                    - **Ngày thực hiện:** 4/16/2024
                    - **CusID:** 22079986
                    - **Tên Khách hàng:** CÔNG TY CỔ PHẦN MẶT DỰNG CAG
                    - **Phân khúc:** MM
                    - **Subsegment:** Dịch vụ Xây lắp, lắp đặt
                    - **XHTD:** Aa3
                    - **BBC:** (Trống)
                    - **DDA:** Vùng
                    </OUTPUT_EXAMPLE>

                    Analyze the provided file. Think step-by-step to ensure all data is captured accurately, then generate the final Markdown output.
                    """


                    response = self.llm_client.generate_content([uploaded_file, prompt])
                    print(f"✅ LLM đã parse thành công file: {original_path.name}")
                    return response.text

                except Exception as e:
                    error_message = str(e)
                    
                    if "429" in error_message or "quota" in error_message.lower():
                        retry_delay = self.base_retry_delay * (2 ** attempt)
                        print(f"  ⚠️ Gặp lỗi Rate Limit. Đang thử lại sau {retry_delay} giây... (Lần {attempt + 1}/{self.max_retries})")
                        time.sleep(retry_delay)
                        continue
                    elif "mimeType" in error_message or "not supported" in error_message:
                        print(f"❌ Lỗi MIME type không được hỗ trợ: {e}")
                        break
                    elif "503" in error_message or "Service Unavailable" in error_message:
                        retry_delay = self.base_retry_delay * (2 ** attempt)
                        print(f"  ⚠️ Gemini service unavailable. Đang thử lại sau {retry_delay} giây... (Lần {attempt + 1}/{self.max_retries})")
                        time.sleep(retry_delay)
                        continue
                    else:
                        print(f"❌ Lỗi khác khi parsing: {e}")
                        break
            
            print(f"  ❌ Đã thử lại {self.max_retries} lần nhưng vẫn gặp lỗi. Bỏ cuộc.")
            return ""
            
        finally:
            # Cleanup
            if uploaded_file:
                try:
                    genai.delete_file(uploaded_file.name)
                    print(f"🧹 Đã cleanup uploaded file: {original_path.name}")
                except:
                    pass
            
            # Remove temporary file if created
            if temp_file_created and file_path.exists():
                try:
                    os.unlink(file_path)
                    print(f"🧹 Đã xóa file tạm: {file_path}")
                except:
                    pass
    
    
    
    # def _parse_with_llm_with_retry(self, file_path: Path) -> str:
    #     """
    #     Parse file với conversion support và retry logic, sử dụng Together AI.
    #     """
    #     original_path = file_path
    #     temp_file_created = False
    #     result_content = "" # Biến để lưu kết quả, thay cho việc return ngay lập tức

    #     # Check if conversion is needed
    #     if file_path.suffix.lower() not in self.supported_direct_types:
    #         if file_path.suffix.lower() in self.convertible_types:
    #             # SỬA LỖI #3: Cập nhật thông báo
    #             print(f"📋 File type {file_path.suffix} không được hỗ trợ trực tiếp. Đang convert...")
    #             try:
    #                 file_path = self._convert_to_supported_format(file_path)
    #                 temp_file_created = True
    #             except Exception as conv_error:
    #                 print(f"❌ Conversion failed: {conv_error}")
    #                 return ""
    #         else:
    #             print(f"❌ File type {file_path.suffix} không được hỗ trợ.")
    #             return ""

    #     print(f"🧠 Bắt đầu parsing bằng LLM (Llama 3.2) cho file: {original_path.name}...")

    #     try:
    #         with open(file_path, "rb") as f:
    #             file_content_base64 = base64.b64encode(f.read()).decode('utf-8')
    #     except Exception as e:
    #         print(f"❌ Lỗi khi đọc và encode file: {file_path.name} - {e}")
    #         return ""

    #     # SỬA LỖI #1: Thêm xử lý cho .txt
    #     mime_type = "application/pdf"
    #     suffix = file_path.suffix.lower()
    #     if suffix in [".png", ".jpg", ".jpeg"]:
    #         mime_type = f"image/{suffix[1:]}"
    #     elif suffix == ".webp":
    #         mime_type = "image/webp"
    #     elif suffix == ".txt":
    #         mime_type = "text/plain"

    #     for attempt in range(self.max_retries):
    #         try:
    #             self._apply_rate_limit()
                
    #             # Sử dụng textwrap.dedent để xóa các khoảng trắng thụt lề
    #             prompt = textwrap.dedent("""
    #                 <TASK_DEFINITION>
    #                 You are an automated data processing engine. Your sole task is to analyze the provided file and convert its entire content into a single, clean, and well-structured Markdown string. The output must be a perfect representation of the original data, suitable for machine parsing later.
    #                 Follow these critical instructions precisely:
    #                 1.  **Analyze Layout:** First, analyze the visual layout of the document. Identify key-value pairs (e.g., a label in one cell and its value in another, potentially non-adjacent cell) and structured tables.
    #                 2.  **Convert Key-Value Pairs:** Represent all identified key-value pairs clearly.
    #                 3.  **Convert Tables:** Convert all structured tables into standard Markdown table format.
    #                 4.  **Preserve Content:** All text and numerical data must be preserved exactly as it appears in the original file.
    #                 5.  **No Extra Content:** Do not add any summaries, explanations, comments, or any text that is not present in the original document.
    #                 6.  **Strict Output Format:** Your entire output must be ONLY the Markdown content. Do not wrap it in ```markdown ... ``` or any other formatting.
    #                 7.  **No Leading/Trailing Whitespace:** The output must not start or end with any whitespace, spaces, or newlines.
    #                 </TASK_DEFINITION>
    #                 <OUTPUT_EXAMPLE>
    #                 ### A. THÔNG TIN CHUNG
    #                 - **Ngày thực hiện:** 4/16/2024
    #                 - **CusID:** 22079986
    #                 - **Tên Khách hàng:** CÔNG TY CỔ PHẦN MẶT DỰNG CAG
    #                 - **Phân khúc:** MM
    #                 - **Subsegment:** Dịch vụ Xây lắp, lắp đặt
    #                 - **XHTD:** Aa3
    #                 - **BBC:** (Trống)
    #                 - **DDA:** Vùng
    #                 </OUTPUT_EXAMPLE>
    #                 Analyze the provided file. Think step-by-step to ensure all data is captured accurately, then generate the final Markdown output.
    #             """)

    #             message = HumanMessage(
    #                 content=[
    #                     {"type": "text", "text": prompt},
    #                     {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{file_content_base64}"}},
    #                 ]
    #             )

    #             response = self.llm_client.invoke([message])
    #             print(f"✅ LLM (Llama 3.2) đã parse thành công file: {original_path.name}")
                
    #             # SỬA LỖI #2: Gán kết quả vào biến và break vòng lặp
    #             raw_content = response.content
    #             result_content = raw_content.strip() # Dọn dẹp khoảng trắng
    #             break  # Thoát khỏi vòng lặp retry khi đã thành công

    #         except Exception as e:
    #             error_message = str(e)
    #             if "429" in error_message or "quota" in error_message.lower():
    #                 retry_delay = self.base_retry_delay * (2 ** attempt)
    #                 print(f"  ⚠️ Gặp lỗi Rate Limit. Đang thử lại sau {retry_delay} giây... (Lần {attempt + 1}/{self.max_retries})")
    #                 time.sleep(retry_delay)
    #                 continue
    #             elif "mimeType" in error_message or "not supported" in error_message:
    #                 print(f"❌ Lỗi MIME type không được hỗ trợ: {e}")
    #                 break
    #             elif "503" in error_message or "Service Unavailable" in error_message:
    #                 retry_delay = self.base_retry_delay * (2 ** attempt)
    #                 # SỬA LỖI #3: Cập nhật thông báo
    #                 print(f"  ⚠️ Service không khả dụng (503). Đang thử lại sau {retry_delay} giây... (Lần {attempt + 1}/{self.max_retries})")
    #                 time.sleep(retry_delay)
    #                 continue
    #             else:
    #                 print(f"❌ Lỗi khác khi parsing: {e}")
    #                 break
        
    #     # Chỉ in thông báo này nếu vòng lặp kết thúc mà không thành công
    #     if not result_content:
    #         print(f"  ❌ Đã thử lại {self.max_retries} lần nhưng vẫn gặp lỗi. Bỏ cuộc.")

    #     # Luôn chạy khối finally để dọn dẹp
    #     try:
    #         if temp_file_created and file_path.exists():
    #             try:
    #                 os.unlink(file_path)
    #                 print(f"🧹 Đã xóa file tạm: {file_path}")
    #             except:
    #                 pass
    #     finally:
    #         # Câu lệnh return duy nhất, giúp code dễ đọc hơn
    #         return result_content

    def parse_file(self, file_path: str) -> List[Document]:
        """
        Parse file với enhanced conversion support.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"❌ File không tồn tại: {file_path}")
            return []
        
        # Parse với LLM
        full_markdown_content = self._parse_with_llm_with_retry(file_path)
        
        if not full_markdown_content or not full_markdown_content.strip():
            print(f"⚠️ LLM không trả về nội dung nào cho file {file_path.name}. Trả về danh sách rỗng.")
            return []
        
        return full_markdown_content
        #Split chunks
        # print(f"🔪 Bắt đầu chunking nội dung Markdown (độ dài: {len(full_markdown_content)} chars)...")
        
        # chunks = self.text_splitter.create_documents([full_markdown_content])
        
        # # Gắn metadata
        # for chunk in chunks:
        #     chunk.metadata = {
        #         "source": file_path.name, # Chỉ lưu tên file cho gọn
        #         "original_path": str(file_path),
        #         "file_type": file_path.suffix.lower(),
        #         "content_type": "llm_parsed_markdown_chunk",
        #         "parser_method": "gemini_enhanced"
        #     }
        
        # print(f"✅ Hoàn tất! Tạo ra {len(chunks)} documents từ file {file_path.name}.")
        # return chunks

    def convert_file_to_pdf(self, file_path: str, output_pdf_path: Optional[str] = None) -> Optional[str]:
        """
        Public method to convert files to PDF format.
        
        Args:
            file_path: Path to input file
            output_pdf_path: Optional output path for PDF
            
        Returns:
            Path to converted PDF file or None if conversion failed
        """
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.convertible_types:
            print(f"❌ Conversion không được hỗ trợ cho file type: {file_extension}")
            return None
        
        try:
            if output_pdf_path:
                output_path = Path(output_pdf_path)
            else:
                output_path = file_path.with_suffix('.pdf')
            
            if file_extension in ['.docx', '.doc']:
                converted_path = self._convert_word_to_pdf_enhanced(file_path)
                print(f"Nội dung của file PDF là : {converted_path}")
            elif file_extension in ['.xlsx', '.xls']:
                converted_path = self._convert_excel_to_pdf_enhanced(file_path)
                print(f"Nội dung của file PDF là : {converted_path}")
            else:
                return None
            
            # Move converted file to desired location if different
            if converted_path != output_path:
                converted_path.rename(output_path)
            
            return str(output_path)
            
        except Exception as e:
            print(f"❌ Conversion failed: {e}")
            return None
        
if __name__ == "__main__":
    parser = DocumentParser()
    parse_folder = Path('/home/locmt/Techcombank_/chatbot_document/data/data_real_new')
    
    # Danh sách để chứa tất cả các chuỗi markdown từ các file
    combined_data = [] 
    
    files_to_parse = list(parse_folder.glob('*'))

    # Lặp qua từng file để parse
    for file_path in files_to_parse:
        if file_path.is_dir():
            continue

        print(f"\n{'='*20} PROCESSING FILE: {file_path.name} {'='*20}")
        
        # Parse file và nhận về một chuỗi Markdown
        markdown_string = parser.parse_file(str(file_path))
        
        if markdown_string:
            print(f"✅ Nhận được chuỗi Markdown từ: {file_path.name}")
            combined_data.append(markdown_string)
        else:
            print(f"❌ Parse thất bại hoặc không có nội dung: {file_path.name}")

    print(f"\n{'='*20} SUMMARY {'='*20}")
    print(f"📄 Đã parse xong {len(combined_data)} file.")
    
    # Gộp tất cả các chuỗi lại và lưu ra file duy nhất (giống hệt notebook)
    output_dir = Path('output_parsing')
    output_dir.mkdir(exist_ok=True)
    output_md_path = output_dir / 'parsed_output.md'

    with open(output_md_path, 'w', encoding='utf-8') as f:
        # Nối tất cả các markdown lại bằng hai dấu xuống dòng
        f.write('\n\n'.join(combined_data))
        
    print(f"✅ Đã gộp và lưu tất cả nội dung vào file: {output_md_path}")        
# if __name__ == "__main__":
#     parser = DocumentParser()
#     parse_folder = Path('/home/locmt/Techcombank_/chatbot_document/data/data_real_new')
    
#     all_parsed_chunks = []  # Danh sách để chứa tất cả các chunk từ tất cả các file
    
#     files_to_parse = list(parse_folder.glob('*'))

#     # Lặp qua từng file để parse
#     for file_path in files_to_parse:
#         if file_path.is_dir():
#             continue

#         print(f"\n{'='*20} PROCESSING FILE: {file_path.name} {'='*20}")
        
#         # Parse file và nhận về một DANH SÁCH CÁC CHUNK
#         document_chunks = parser.parse_file(str(file_path))
        
#         if document_chunks:
#             print(f"✅ Parse thành công: {file_path.name}, tạo ra {len(document_chunks)} chunks.")
#             all_parsed_chunks.extend(document_chunks)
#         else:
#             print(f"❌ Parse thất bại hoặc không có nội dung: {file_path.name}")

#     print(f"\n{'='*20} SUMMARY {'='*20}")
#     print(f"📄 Đã parse xong toàn bộ thư mục. Tổng cộng có {len(all_parsed_chunks)} chunks.")
    
#     # Lưu kết quả ra một file JSON duy nhất chứa tất cả các chunk
#     output_dir = Path('output_parsing')
#     output_dir.mkdir(exist_ok=True)
#     output_json_path = output_dir / 'all_chunks.json'

#     # Chuyển đổi Document objects thành dict để có thể serialize
#     chunks_as_dicts = [
#         {"page_content": doc.page_content, "metadata": doc.metadata} 
#         for doc in all_parsed_chunks
#     ]

#     with open(output_json_path, 'w', encoding='utf-8') as f:
#         json.dump(chunks_as_dicts, f, indent=2, ensure_ascii=False)
        
#     print(f"✅ Đã lưu tất cả các chunk vào file JSON: {output_json_path}")